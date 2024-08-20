from docx import Document
from docx.shared import Inches

def insert_image_after_text(doc, text, image_path):
    for paragraph in doc.paragraphs:
        if text in paragraph.text:
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(6))
            break

def add_images(docx_path, text_image_pairs):
    doc = Document(docx_path)
    for text, image_path in text_image_pairs:
        insert_image_after_text(doc, text, image_path)
    doc.save(docx_path)

if __name__ == "__main__":
    docx_path = 'AutomatizacionSuspencionReconexion/Suspencion.docx'
    text_image_pairs = [

        ("Se busca el cliente sobre el cual se va a aplicar la suspencion", 'AutomatizacionSuspencionReconexion/Evidencia_SUSPENCIONscreenshots/captura1.png'),

        ("Se valida los productos instalados", 'AutomatizacionSuspencionReconexion/Evidencia_SUSPENCIONscreenshots/captura2.png'),

        ("Se verifica el estado completado en vista de detalles", 'AutomatizacionSuspencionReconexion/Evidencia_SUSPENCIONscreenshots/captura3.png'),

        ("Confirmamos en productos instalados los estados en Suspendido", 'AutomatizacionSuspencionReconexion/Evidencia_SUSPENCIONscreenshots/captura4.png'),

        
    ]

    add_images(docx_path, text_image_pairs)
