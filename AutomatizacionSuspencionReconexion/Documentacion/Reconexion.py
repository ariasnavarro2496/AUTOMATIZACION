from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_font_style(run, font_name='Calibri', font_size=10):
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # For compatibility with older versions of Word
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

def crear_documento(ruta_guardado):
    # Crear un nuevo documento
    doc = Document()

    # Agregar un título
    title = doc.add_heading('Evidencia Reconexion', level=1)
    set_font_style(title.runs[0])

    # Agregar otro párrafo con texto en negrita
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Descripción:')
    run.bold = True
    set_font_style(run)

    #Agregar otro párrafo con texto en cursiva
    run = paragraph.add_run(' Reconexion.')
    run.italic = True
    set_font_style(run)

    # Agregar un párrafo
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Se busca el cliente sobre el cual se va a aplicar la Reconexion')
    set_font_style(run)

    # Agregar un párrafo
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Se valida los productos Suspendidos a los cual se les aplicara Reconexion')
    set_font_style(run)

    # Agregar un párrafo
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Se verifica el estado completado en vista de detalles ')
    set_font_style(run)

    # Agregar un párrafo
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Confirmamos en productos instalados los estados en Activo nuevamente')
    set_font_style(run)



    # Crear el directorio si no existe
    os.makedirs(os.path.dirname(ruta_guardado), exist_ok=True)

    # Guardar el documento
    doc.save(ruta_guardado)

# Ejemplo de uso:
if __name__ == "__main__":
    # Definir la ruta relativa donde se guardará el documento dentro del proyecto Robot Framework
    ruta_guardado = os.path.join('AutomatizacionSuspencionReconexion', 'Reconexion.docx')
    crear_documento(ruta_guardado)
