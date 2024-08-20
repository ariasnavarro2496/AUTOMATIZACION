from docx import Document
from docx.shared import Inches

def insert_image_after_text(doc, text, image_path):
    for paragraph in doc.paragraphs:
        if text in paragraph.text:
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(6))
            break

def add_images(docx_path, text_image_pairs):
    doc = Document(docx_path)
    for text, image_path in text_image_pairs:
        insert_image_after_text(doc, text, image_path)
    doc.save(docx_path)

if __name__ == "__main__":
    docx_path = 'AutomatizacionALTASCLARO/ALTA3PLAY_CLARO.docx'
    text_image_pairs = [
        ("Se Añade dirección de la promoción y se suma el alta play 3", 'AutomatizacionALTASCLARO/Evidencia_CLAROscreenshots/captura1.png'),

        ("Se llena los campos de cliente nuevo y genera el RUT", 'AutomatizacionALTASCLARO/Evidencia_CLAROscreenshots/captura2.png'),

        ("Selecciona la promocion", 'AutomatizacionALTASCLARO/Evidencia_CLAROscreenshots/captura3.png'),

        ("A nivel de la orden, solicitar TSQ, Generar documento y se Envía el pedido", 'AutomatizacionALTASCLARO/Evidencia_CLAROscreenshots/captura4.png'),

        ("Se genera la actividad", 'AutomatizacionALTASCLARO/Evidencia_CLAROscreenshots/captura5.png'),

        ("Instalamos equipo y aprovisionamos", 'AutomatizacionALTASCLARO/Evidencia_CLAROscreenshots/captura8.png'),

        ("Finalizamos actividad", 'AutomatizacionALTASCLARO/Evidencia_CLAROscreenshots/captura9.png'),
        
        ("En Siebel se revisa que la actividad aparezca completada ", 'AutomatizacionALTASCLARO/Evidencia_CLAROscreenshots/captura10.png')
    ]

    add_images(docx_path, text_image_pairs)
